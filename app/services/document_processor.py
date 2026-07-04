"""Document processing service for Spaces.

Handles extraction of text from various file formats:
- Word documents (.docx)
- Excel spreadsheets (.xlsx)
- PDF files (.pdf)
- Web links
- YouTube video transcripts
- Audio/voice transcription via Whisper
- Image description via GPT-4V
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional, Tuple
from pathlib import Path
from urllib.parse import urlparse, parse_qs

from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


# ============================================================
# YOUTUBE HELPERS
# ============================================================


def extract_youtube_video_id(url: str) -> Optional[str]:
    """
    Extract video ID from various YouTube URL formats:
    - https://www.youtube.com/watch?v=VIDEO_ID
    - https://youtu.be/VIDEO_ID
    - https://www.youtube.com/embed/VIDEO_ID
    - https://www.youtube.com/v/VIDEO_ID
    - https://www.youtube.com/shorts/VIDEO_ID
    """
    if not url:
        return None

    parsed = urlparse(url)

    # youtu.be short links
    if parsed.netloc in ('youtu.be', 'www.youtu.be'):
        return parsed.path.lstrip('/')

    # youtube.com links
    if parsed.netloc in ('youtube.com', 'www.youtube.com', 'm.youtube.com'):
        # /watch?v=VIDEO_ID
        if parsed.path == '/watch':
            query = parse_qs(parsed.query)
            return query.get('v', [None])[0]

        # /embed/VIDEO_ID, /v/VIDEO_ID, /shorts/VIDEO_ID
        for prefix in ('/embed/', '/v/', '/shorts/'):
            if parsed.path.startswith(prefix):
                return parsed.path[len(prefix):].split('?')[0].split('&')[0]

    return None


def is_youtube_url(url: str) -> bool:
    """Check if URL is a YouTube link."""
    return extract_youtube_video_id(url) is not None


# ============================================================
# LANGUAGE DETECTION
# ============================================================


def detect_language(text: str) -> str:
    """
    Simple heuristic language detection.
    Returns 'ru' for Russian, 'en' for English, 'uk' for Ukrainian, etc.
    """
    if not text:
        return "ru"

    # Count Cyrillic vs Latin characters
    cyrillic_count = len(re.findall(r'[а-яА-ЯёЁ]', text))
    latin_count = len(re.findall(r'[a-zA-Z]', text))

    # Ukrainian-specific characters
    ukrainian_chars = len(re.findall(r'[іїєґІЇЄҐ]', text))

    if cyrillic_count > latin_count:
        if ukrainian_chars > cyrillic_count * 0.05:  # >5% Ukrainian chars
            return "uk"
        return "ru"
    elif latin_count > 0:
        return "en"

    return "ru"


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================


async def extract_word(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """
    Extract text from Word document (.docx).
    Returns: (extracted_text, error_message)
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)
        return text, None

    except ImportError:
        return "", "Библиотека python-docx не установлена"
    except Exception as e:
        logger.error(f"Error extracting Word document: {e}")
        return "", f"Ошибка обработки Word: {str(e)}"


async def extract_excel(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """
    Extract text from Excel spreadsheet (.xlsx).
    Returns: (extracted_text, error_message)
    """
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        all_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            all_text.append(f"=== Лист: {sheet_name} ===")

            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))

                if row_values:
                    all_text.append(" | ".join(row_values))

        text = "\n".join(all_text)
        return text, None

    except ImportError:
        return "", "Библиотека openpyxl не установлена"
    except Exception as e:
        logger.error(f"Error extracting Excel document: {e}")
        return "", f"Ошибка обработки Excel: {str(e)}"


async def extract_pdf(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """
    Extract text from PDF file.
    Returns: (extracted_text, error_message)
    """
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        all_text = []

        for i, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                all_text.append(f"--- Страница {i} ---\n{page_text.strip()}")

        text = "\n\n".join(all_text)
        return text, None

    except ImportError:
        return "", "Библиотека PyPDF2 не установлена"
    except Exception as e:
        logger.error(f"Error extracting PDF document: {e}")
        return "", f"Ошибка обработки PDF: {str(e)}"


# ============================================================
# YOUTUBE TRANSCRIPT
# ============================================================


def _fetch_youtube_transcript_sync(video_id: str) -> Tuple[str, Optional[str]]:
    """
    Synchronous function to fetch YouTube transcript.
    Called via run_in_executor for async compatibility.
    """
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        from youtube_transcript_api._errors import (
            TranscriptsDisabled,
            NoTranscriptFound,
            VideoUnavailable,
        )

        ytt_api = YouTubeTranscriptApi()

        # Try to get transcript in preferred languages
        # Priority: Russian, then English, then any available
        transcript = None
        video_title = f"YouTube видео (ID: {video_id})"

        try:
            # First try Russian subtitles
            transcript = ytt_api.fetch(video_id, languages=['ru'])
        except NoTranscriptFound:
            try:
                # Try English
                transcript = ytt_api.fetch(video_id, languages=['en'])
            except NoTranscriptFound:
                try:
                    # Try any available transcript
                    transcript_list = ytt_api.list(video_id)
                    # Get first available
                    for t in transcript_list:
                        transcript = t.fetch()
                        break
                except Exception:
                    pass

        if not transcript:
            return "", "Субтитры для этого видео недоступны"

        # Combine transcript segments into text
        # Note: segments are FetchedTranscriptSnippet objects with .text attribute
        text_parts = []
        for segment in transcript:
            # Handle both old dict format and new object format
            if hasattr(segment, 'text'):
                text_parts.append(segment.text)
            elif isinstance(segment, dict):
                text_parts.append(segment.get('text', ''))

        full_text = ' '.join(text_parts)

        # Clean up common transcript artifacts
        full_text = re.sub(r'\[.*?\]', '', full_text)  # Remove [Music], [Applause], etc.
        full_text = re.sub(r'\s+', ' ', full_text).strip()

        result = f"📺 {video_title}\n\n{full_text}"
        return result[:20000], None  # Limit to 20k chars

    except TranscriptsDisabled:
        return "", "Субтитры отключены для этого видео"
    except VideoUnavailable:
        return "", "Видео недоступно"
    except ImportError:
        return "", "Библиотека youtube-transcript-api не установлена"
    except Exception as e:
        error_str = str(e).lower()
        logger.error(f"Error fetching YouTube transcript: {e}")

        # User-friendly error messages
        if "age-restricted" in error_str or "age restricted" in error_str:
            return "", (
                "🔞 Видео имеет возрастное ограничение (18+).\n"
                "Субтитры недоступны без авторизации YouTube.\n"
                "Попробуйте другое видео или обратитесь к администратору."
            )
        elif "ip" in error_str and ("ban" in error_str or "block" in error_str):
            return "", (
                "🚫 IP-адрес сервера заблокирован YouTube.\n"
                "Обратитесь к администратору для решения проблемы."
            )
        elif "403" in error_str or "forbidden" in error_str:
            return "", (
                "⛔ Доступ запрещён YouTube.\n"
                "Возможные причины: возрастное ограничение, гео-блокировка или лимит запросов.\n"
                "Попробуйте позже или обратитесь к администратору."
            )
        elif "404" in error_str or "not found" in error_str:
            return "", "Видео не найдено. Проверьте ссылку."
        elif "private" in error_str:
            return "", "Это приватное видео. Субтитры недоступны."
        else:
            return "", f"Ошибка получения субтитров: {str(e)[:200]}"


async def fetch_youtube_transcript(url: str) -> Tuple[str, Optional[str]]:
    """
    Fetch transcript/subtitles from YouTube video.
    Returns: (transcript_text, error_message)
    """
    import asyncio

    video_id = extract_youtube_video_id(url)
    if not video_id:
        return "", "Не удалось извлечь ID видео из ссылки"

    # Run sync function in executor to avoid blocking event loop
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _fetch_youtube_transcript_sync, video_id)


# ============================================================
# WEB LINK EXTRACTION
# ============================================================


def _fetch_link_content_sync(url: str, timeout: int = 30) -> Tuple[str, Optional[str]]:
    """
    Synchronous function to fetch web content.
    Called via run_in_executor for async compatibility.
    """
    import requests

    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
        }

        response = requests.get(url, headers=headers, timeout=timeout, verify=True)

        if response.status_code != 200:
            return "", f"HTTP ошибка: {response.status_code}"

        html = response.text

        soup = BeautifulSoup(html, "lxml")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
            element.decompose()

        # Try to find main content
        main_content = None
        for selector in ["article", "main", ".content", ".post", "#content", ".article-body"]:
            main_content = soup.select_one(selector)
            if main_content:
                break

        if main_content:
            text = main_content.get_text(separator="\n", strip=True)
        else:
            # Fallback to body
            body = soup.find("body")
            text = body.get_text(separator="\n", strip=True) if body else soup.get_text(separator="\n", strip=True)

        # Clean up excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Get title
        title = soup.find("title")
        title_text = title.get_text(strip=True) if title else ""

        if title_text:
            text = f"🔗 {title_text}\n\n{text}"

        return text[:15000], None  # Limit to 15k chars

    except requests.exceptions.Timeout:
        return "", "Таймаут при загрузке страницы"
    except requests.exceptions.RequestException as e:
        logger.error(f"Error fetching URL {url}: {e}")
        return "", f"Ошибка загрузки: {str(e)}"
    except Exception as e:
        logger.error(f"Error processing URL {url}: {e}")
        return "", f"Ошибка обработки ссылки: {str(e)}"


async def fetch_link_content(url: str, timeout: int = 30) -> Tuple[str, Optional[str]]:
    """
    Fetch and extract main content from a web URL.
    Returns: (extracted_text, error_message)
    """
    import asyncio

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _fetch_link_content_sync, url, timeout)


# ============================================================
# AUDIO TRANSCRIPTION
# ============================================================


async def transcribe_audio(
    file_bytes: bytes,
    file_name: str,
    openai_client,
) -> Tuple[str, Optional[str]]:
    """
    Transcribe audio using OpenAI Whisper.
    Returns: (transcribed_text, error_message)
    """
    try:
        # Determine file extension
        suffix = Path(file_name).suffix.lower() or ".ogg"

        # Create file-like object with proper name
        audio_file = io.BytesIO(file_bytes)
        audio_file.name = f"audio{suffix}"

        response = await openai_client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            language="ru",  # Default to Russian
        )

        return response.text, None

    except Exception as e:
        logger.error(f"Error transcribing audio: {e}")
        return "", f"Ошибка транскрипции: {str(e)}"


# ============================================================
# IMAGE DESCRIPTION
# ============================================================


async def describe_image(
    file_bytes: bytes,
    openai_client,
) -> Tuple[str, Optional[str]]:
    """
    Describe image content using GPT-4V.
    Returns: (description_text, error_message)
    """
    try:
        import base64

        # Encode image to base64
        image_base64 = base64.b64encode(file_bytes).decode("utf-8")

        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Опиши подробно что изображено на этой картинке. "
                                   "Опиши текст, если он есть. Ответ на русском языке."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_base64}",
                                "detail": "low"
                            }
                        }
                    ]
                }
            ],
            max_tokens=500,
        )

        text = response.choices[0].message.content
        return text, None

    except Exception as e:
        logger.error(f"Error describing image: {e}")
        return "", f"Ошибка описания изображения: {str(e)}"


# ============================================================
# TRANSLATION
# ============================================================


async def translate_to_russian(
    text: str,
    source_language: str,
    openai_client,
) -> Tuple[str, Optional[str]]:
    """
    Translate text to Russian using LLM.
    Returns: (translated_text, error_message)
    """
    if source_language == "ru":
        return text, None

    if not text or len(text.strip()) < 10:
        return text, None

    try:
        lang_names = {
            "en": "английского",
            "uk": "украинского",
            "de": "немецкого",
            "fr": "французского",
            "es": "испанского",
        }
        lang_name = lang_names.get(source_language, "иностранного")

        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"Переведи текст с {lang_name} на русский язык. "
                               "Сохрани форматирование и структуру. "
                               "Отвечай только переводом, без комментариев."
                },
                {
                    "role": "user",
                    "content": text[:8000]  # Limit input
                }
            ],
            max_tokens=4000,
        )

        translated = response.choices[0].message.content
        return translated, None

    except Exception as e:
        logger.error(f"Error translating text: {e}")
        return text, f"Ошибка перевода: {str(e)}"


# ============================================================
# MAIN PROCESSOR
# ============================================================


async def process_material(
    material_type: str,
    file_bytes: Optional[bytes],
    file_name: Optional[str],
    content: Optional[str],
    openai_client,
    auto_translate: bool = True,
) -> Tuple[str, str, Optional[str]]:
    """
    Process a material based on its type.

    Returns: (processed_text, detected_language, error_message)
    """
    processed_text = ""
    error = None

    if material_type == "text":
        processed_text = content or ""

    elif material_type == "voice" or material_type == "audio":
        if file_bytes:
            processed_text, error = await transcribe_audio(
                file_bytes, file_name or "audio.ogg", openai_client
            )

    elif material_type == "image":
        if file_bytes:
            processed_text, error = await describe_image(file_bytes, openai_client)

    elif material_type == "document_word":
        if file_bytes:
            processed_text, error = await extract_word(file_bytes)

    elif material_type == "document_excel":
        if file_bytes:
            processed_text, error = await extract_excel(file_bytes)

    elif material_type == "document_pdf":
        if file_bytes:
            processed_text, error = await extract_pdf(file_bytes)

    elif material_type == "link":
        if content:
            processed_text, error = await fetch_link_content(content)

    elif material_type == "youtube":
        if content:
            processed_text, error = await fetch_youtube_transcript(content)

    elif material_type == "video":
        # Video processing not implemented yet
        processed_text = "[Видео - обработка недоступна]"

    # Detect language
    language = detect_language(processed_text)

    # Auto-translate if needed
    if auto_translate and language != "ru" and processed_text and not error:
        translated, translate_error = await translate_to_russian(
            processed_text, language, openai_client
        )
        if not translate_error:
            processed_text = f"[Оригинал: {language}]\n\n{translated}"

    return processed_text, language, error
